from pathlib import Path

from pdftransl.export.exporter import available_engines, export_document
from pdftransl.export.html import markdown_to_html

MD = """# Заголовок статьи

Абзац с **жирным**, *курсивом* и формулой $E = mc^2$.

$$
\\int_0^1 x\\,dx = \\frac{1}{2}
$$

| Модель | Точность |
|---|---|
| CNN | 0.92 |

![Рисунок 1](images/fig1.png)

```python
print("код")
```
"""


def test_html_export_structure(tmp_path):
    html = markdown_to_html(MD, title="Тест")
    assert "<h1>Заголовок статьи</h1>" in html
    assert "<strong>жирным</strong>" in html
    assert "$E = mc^2$" in html                 # formula left for KaTeX
    assert "\\int_0^1" in html
    assert "<table>" in html and "<th>Модель</th>" in html
    assert "katex" in html                      # renderer wired
    assert "<pre><code>" in html


def test_html_inlines_images(tmp_path):
    assets = tmp_path / "assets" / "images"
    assets.mkdir(parents=True)
    # 1x1 png
    png = bytes.fromhex(
        "89504e470d0a1a0a0000000d49484452000000010000000108060000001f15c489"
        "0000000d49444154789c626001000000ffff03000006000557bfabd40000000049454e44ae426082"
    )
    (assets / "fig1.png").write_bytes(png)
    html = markdown_to_html(MD, assets_dir=tmp_path / "assets")
    assert "data:image/png;base64," in html


def test_export_docx_native(tmp_path):
    result = export_document(MD, tmp_path / "doc", formats=["docx"])
    docx_path = result["files"]["docx"]
    assert docx_path is not None and Path(docx_path).exists()
    assert result["engines"]["docx"] in ("pandoc", "python-docx")

    import docx

    document = docx.Document(docx_path)
    texts = [p.text for p in document.paragraphs]
    assert any("Заголовок статьи" in t for t in texts)
    assert any("E = mc^2" in t for t in texts)   # formula kept as LaTeX text
    assert len(document.tables) == 1
    assert document.tables[0].cell(0, 0).text == "Модель"


def test_export_html_and_missing_pdf_engine_reported(tmp_path):
    result = export_document(MD, tmp_path / "doc", formats=["html", "pdf"])
    assert Path(result["files"]["html"]).exists()
    # pdf may or may not be exportable in this environment; either a file
    # exists or the reason is reported
    if result["files"]["pdf"] is None:
        assert "unavailable" in result["engines"]["pdf"]
    else:
        assert Path(result["files"]["pdf"]).exists()


def test_available_engines_shape():
    engines = available_engines()
    assert engines["html"] == ["builtin"]
    assert isinstance(engines["docx"], list)
    assert isinstance(engines["pdf"], list)

"""Tests for the export-formula and scanned-PDF fixes."""

import io

import pytest

from pdftransl.config import PipelineConfig
from pdftransl.export.formula_render import matplotlib_available, render_latex_png, strip_math_delimiters
from pdftransl.export.katex_assets import find_katex_dist, is_offline_capable, katex_head
from pdftransl.llm.fake import FakeLLMClient

fitz = pytest.importorskip("fitz", reason="PyMuPDF needed for scan tests")


# ---- helpers ------------------------------------------------------------

def _text_pdf(path):
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 100), "A Study of Gradient Descent", fontsize=16)
    page.insert_text((72, 130), "We analyze convergence of stochastic descent.", fontsize=11)
    doc.save(str(path))
    doc.close()


def _scanned_pdf(path):
    """A PDF whose only content is a rasterized page (no text layer)."""
    src = fitz.open()
    page = src.new_page()
    page.insert_text((72, 100), "Scanned title here", fontsize=16)
    page.insert_text((72, 130), "Some scanned body text.", fontsize=11)
    pix = page.get_pixmap(dpi=120)
    scan = fitz.open()
    p = scan.new_page(width=page.rect.width, height=page.rect.height)
    p.insert_image(p.rect, pixmap=pix)
    scan.save(str(path))
    scan.close()
    src.close()


# ---- scan detection ------------------------------------------------------

def test_scan_detection(tmp_path):
    from pdftransl.parsing.scan_detect import scan_stats

    text_pdf = tmp_path / "text.pdf"
    scan_pdf = tmp_path / "scan.pdf"
    _text_pdf(text_pdf)
    _scanned_pdf(scan_pdf)

    assert scan_stats(text_pdf)["is_scanned"] is False
    stats = scan_stats(scan_pdf)
    assert stats["is_scanned"] is True
    assert stats["image_pages"] == 1


# ---- VLM OCR backend -----------------------------------------------------

class FakeVision(FakeLLMClient):
    supports_vision = True

    def chat(self, messages, temperature=0.2, max_tokens=None, response_format=None):
        self.calls.append(messages)
        content = messages[-1]["content"]
        # OCR requests must carry the page image
        assert isinstance(content, list)
        assert any(p.get("type") == "image_url" for p in content)
        return "# Scanned title\n\nBody with formula $x^2$."


def test_vlm_ocr_backend(tmp_path):
    from pdftransl.parsing.vlm_ocr_backend import VlmOcrBackend

    scan_pdf = tmp_path / "scan.pdf"
    _scanned_pdf(scan_pdf)
    backend = VlmOcrBackend(PipelineConfig(ocr_dpi=120), client=FakeVision())
    assert backend.available() is True
    parsed = backend.parse(scan_pdf, tmp_path / "ocr")
    assert parsed.backend == "vlm_ocr"
    assert "$x^2$" in parsed.markdown
    assert parsed.meta["ocr"] is True
    assert parsed.assets and parsed.assets[0].kind == "page"


class FakeVisionAndText(FakeLLMClient):
    """Vision-capable client used for both OCR (image) and translation (text)."""

    supports_vision = True

    def chat(self, messages, temperature=0.2, max_tokens=None, response_format=None):
        self.calls.append(messages)
        content = messages[-1]["content"]
        if isinstance(content, list):  # OCR request with a page image
            return "# Scanned title\n\nBody with formula $x^2$ here."
        return content if isinstance(content, str) else str(content)  # echo translate


def test_pipeline_routes_scan_to_ocr(tmp_path):
    from pdftransl.pipeline import TranslationPipeline

    scan_pdf = tmp_path / "scan.pdf"
    _scanned_pdf(scan_pdf)
    cfg = PipelineConfig(
        parser_backend="pymupdf", db_path=str(tmp_path / "db.sqlite"),
        output_dir=str(tmp_path / "out"), review=False, doc_summary=False,
        auto_glossary=False, learn=False, export_formats=[], ocr_dpi=120,
    )
    pipeline = TranslationPipeline(cfg, client=FakeVisionAndText())
    result = pipeline.run(scan_pdf, job_id="s1")
    assert result.report["parser_backend"] == "vlm_ocr"
    assert result.report["scan"]["is_scanned"] is True
    assert result.report.get("ocr")


def test_pipeline_warns_when_no_vision(tmp_path):
    from pdftransl.pipeline import TranslationPipeline

    scan_pdf = tmp_path / "scan.pdf"
    _scanned_pdf(scan_pdf)
    text_only = FakeLLMClient(transform=lambda t: t)  # supports_vision = False
    cfg = PipelineConfig(
        parser_backend="pymupdf", db_path=str(tmp_path / "db.sqlite"),
        output_dir=str(tmp_path / "out"), review=False, doc_summary=False,
        auto_glossary=False, learn=False, export_formats=[],
    )
    pipeline = TranslationPipeline(cfg, client=text_only)
    result = pipeline.run(scan_pdf, job_id="s2")
    assert result.report["parser_backend"] == "pymupdf"
    assert "scan_warning" in result.report


# ---- KaTeX vendoring -----------------------------------------------------

def test_katex_offline_when_vendored():
    dist = find_katex_dist()
    if dist is None:
        pytest.skip("KaTeX not vendored (run npm install in frontend/)")
    assert is_offline_capable()
    head = katex_head(offline=True)
    assert "cdn.jsdelivr" not in head
    assert "data:font/woff2;base64," in head
    assert "url(fonts/" not in head          # no broken relative refs


def test_katex_falls_back_to_cdn(monkeypatch):
    import pdftransl.export.katex_assets as ka

    monkeypatch.setattr(ka, "find_katex_dist", lambda: None)
    head = ka.katex_head(offline=True)
    assert "cdn.jsdelivr" in head


def test_html_export_self_contained_offline():
    from pdftransl.export.html import markdown_to_html

    if not is_offline_capable():
        pytest.skip("KaTeX not vendored")
    html = markdown_to_html("# T\n\n$$E=mc^2$$\n", offline=True)
    assert "cdn.jsdelivr" not in html


# ---- DOCX formula rendering ---------------------------------------------

def test_render_latex_png_simple():
    if not matplotlib_available():
        pytest.skip("matplotlib not installed")
    png = render_latex_png("E = mc^2")
    assert png and png[:8] == b"\x89PNG\r\n\x1a\n"


def test_render_latex_png_rejects_unsupported():
    if not matplotlib_available():
        pytest.skip("matplotlib not installed")
    # aligned environment is outside mathtext's subset -> None (caller keeps text)
    assert render_latex_png(r"\begin{aligned} a &= b \\ c &= d \end{aligned}") is None


def test_strip_math_delimiters():
    assert strip_math_delimiters("$$E=mc^2$$") == "E=mc^2"
    assert strip_math_delimiters("$x$") == "x"
    assert strip_math_delimiters("\\[a\\]") == "a"


def test_docx_embeds_formula_images(tmp_path):
    pytest.importorskip("docx")
    if not matplotlib_available():
        pytest.skip("matplotlib not installed")
    from pdftransl.export.docx_native import export_docx

    md = "# T\n\nInline $E=mc^2$ formula.\n\n$$\n\\frac{a}{b}\n$$\n"
    out = export_docx(md, tmp_path / "d.docx")
    import docx

    document = docx.Document(str(out))
    assert len(document.inline_shapes) == 2   # inline + display rendered as images


def test_docx_sanitizes_control_chars(tmp_path):
    pytest.importorskip("docx")
    from pdftransl.export.docx_native import export_docx

    # a stray control character must not crash the writer
    out = export_docx("# Title\x0c\n\nBody\x07 text.\n", tmp_path / "c.docx")
    import docx

    document = docx.Document(str(out))
    assert any("Body" in p.text for p in document.paragraphs)

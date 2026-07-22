"""DOCX без pandoc: python-docx.

Формулы рендерятся картинками через matplotlib mathtext — в документе
видна настоящая формула, а не сырой $...$; что mathtext не осилил —
падает обратно в текст.
"""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Optional

from pdftransl.export.formula_render import render_latex_png, strip_math_delimiters
from pdftransl.models import BlockType
from pdftransl.parsing.splitter import split_markdown

_INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*|\*[^*\n]+\*|`[^`]+`|\$[^$\n]+\$)"
)

# XML/Word reject control characters; strip everything except \t and \n
_CONTROL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


def _clean(text: str) -> str:
    return _CONTROL_RE.sub("", text)


def _add_inline_formula(paragraph, formula: str) -> bool:
    """Add an inline formula as a small image; return True on success."""
    png = render_latex_png(strip_math_delimiters(formula), fontsize=12, dpi=200)
    if png is None:
        return False
    try:
        from docx.shared import Pt

        run = paragraph.add_run()
        # height tuned to sit on the text baseline (~ font size)
        run.add_picture(io.BytesIO(png), height=Pt(13))
        return True
    except Exception:
        return False


def _add_runs(paragraph, text: str) -> None:
    """Split inline markdown into styled runs."""
    for part in _INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        elif part.startswith("$") and part.endswith("$"):
            if not _add_inline_formula(paragraph, part):
                run = paragraph.add_run(_clean(part))   # fallback: LaTeX text
                run.font.name = "Cambria Math"
        else:
            # strip residual link/image syntax to plain text
            clean = re.sub(r"!?\[([^\]]*)\]\([^)]*\)", r"\1", part)
            paragraph.add_run(_clean(clean))


def _find_image(src: str, assets_dir: Optional[Path]) -> Optional[Path]:
    if assets_dir is None:
        return None
    for candidate in (assets_dir / src, assets_dir / Path(src).name):
        if candidate.exists():
            return candidate
    return None


def export_docx(
    markdown: str,
    output_path: str | Path,
    assets_dir: Optional[str | Path] = None,
    title: str = "",
) -> Path:
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    assets = Path(assets_dir) if assets_dir else None

    document = docx.Document()
    if title:
        document.core_properties.title = title

    for block in split_markdown(markdown):
        text = block.text
        if block.type == BlockType.HEADING:
            match = re.match(r"^(#{1,6})\s*(.*)$", text)
            if not match:
                document.add_heading(_clean(text), level=1)
                continue
            document.add_heading(_clean(match.group(2)), level=min(len(match.group(1)), 9))
        elif block.type == BlockType.MATH:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            png = render_latex_png(strip_math_delimiters(text), fontsize=15, dpi=200)
            if png is not None:
                try:
                    # natural size from the PNG's embedded DPI, capped at 6"
                    pic = paragraph.add_run().add_picture(io.BytesIO(png))
                    if pic.width > Inches(6):
                        pic.height = int(pic.height * Inches(6) / pic.width)
                        pic.width = Inches(6)
                except Exception:
                    png = None
            if png is None:  # fallback: LaTeX text
                run = paragraph.add_run(_clean(text))
                run.font.name = "Cambria Math"
                run.font.size = Pt(10)
        elif block.type == BlockType.CODE:
            body = re.sub(r"^```[^\n]*\n?|\n?```$", "", text)
            paragraph = document.add_paragraph()
            run = paragraph.add_run(_clean(body))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif block.type == BlockType.TABLE:
            rows = [r for r in text.splitlines() if r.strip().startswith("|")]
            grid = []
            for row in rows:
                cells = [c.strip() for c in row.strip().strip("|").split("|")]
                if all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c):
                    continue
                grid.append(cells)
            if grid:
                cols = max(len(r) for r in grid)
                table = document.add_table(rows=len(grid), cols=cols)
                table.style = "Light Grid Accent 1"
                for i, row_cells in enumerate(grid):
                    for j, cell in enumerate(row_cells):
                        _add_runs(table.cell(i, j).paragraphs[0], cell)
        elif block.type == BlockType.IMAGE:
            match = re.match(r"^\s*!\[([^\]]*)\]\(([^)]*)\)\s*$", text)
            if not match:
                document.add_paragraph(_clean(text))
                continue
            path = _find_image(match.group(2), assets)
            if path is not None:
                try:
                    document.add_picture(str(path), width=Inches(5.5))
                    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    document.add_paragraph(f"[figure: {match.group(2)}]")
            else:
                document.add_paragraph(f"[figure: {match.group(2)}]")
            if match.group(1):
                caption = document.add_paragraph(match.group(1))
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.runs[0].italic = True
        elif block.type == BlockType.HTML:
            continue  # raw html has no docx representation
        else:
            paragraph = document.add_paragraph()
            _add_runs(paragraph, text.replace("\n", " "))

    document.save(str(output_path))
    return output_path
